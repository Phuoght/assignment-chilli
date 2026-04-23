from __future__ import annotations

import argparse
import csv
import json
import logging
import os
import re
import sys
import textwrap
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Generator, Iterable

logging.basicConfig(level=logging.INFO, format="%(levelname)s │ %(message)s")
log = logging.getLogger("chunker")

# ──────────────────────────────────────────────────────────────────────────────
# DATA MODEL
# ──────────────────────────────────────────────────────────────────────────────

@dataclass
class Chunk:
    text: str
    metadata: dict = field(default_factory=dict)

    def to_jsonl(self) -> str:
        return json.dumps({"text": self.text, "metadata": self.metadata},
                          ensure_ascii=False)


# ──────────────────────────────────────────────────────────────────────────────
# UTILITIES
# ──────────────────────────────────────────────────────────────────────────────

def _slug(path: Path) -> str:
    return path.stem.replace("_", " ").replace("-", " ").title()


def _fixed_windows(text: str, size: int, overlap: int) -> list[str]:
    """Split text into overlapping fixed-size character windows."""
    if len(text) <= size:
        return [text]
    chunks, start = [], 0
    while start < len(text):
        end = start + size
        chunks.append(text[start:end])
        start += size - overlap
    return chunks


def _clean(text: str) -> str:
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _is_heading_style(style_name: str) -> bool:
    return style_name.lower().startswith("heading")


# ──────────────────────────────────────────────────────────────────────────────
# EXTRACTORS  (text + optional pre-split by section)
# ──────────────────────────────────────────────────────────────────────────────

def extract_docx(path: Path) -> list[dict]:
    """
    Returns a list of blocks: {"text": str, "section": str, "type": str}
    Handles paragraphs, tables, and lists.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        log.error("python-docx not installed: pip install python-docx")
        return []

    doc = Document(path)
    blocks = []
    current_section = "Document"

    for para in doc.paragraphs:
        text = _clean(para.text)
        if not text:
            continue
        style = para.style.name if para.style else ""

        if _is_heading_style(style):
            current_section = text
            blocks.append({"text": text, "section": current_section, "type": "heading"})
        elif style.lower().startswith("list"):
            blocks.append({"text": f"• {text}", "section": current_section, "type": "list"})
        else:
            blocks.append({"text": text, "section": current_section, "type": "paragraph"})

    # Extract tables as structured text
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [_clean(c.text) for c in row.cells if _clean(c.text)]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            table_text = "\n".join(rows)
            blocks.append({"text": table_text, "section": current_section, "type": "table"})

    return blocks


def extract_pdf(path: Path) -> list[dict]:
    """
    Tries pdfplumber first (better layout), falls back to PyPDF2.
    Returns blocks with page-level section context.
    """
    blocks = []

    # --- pdfplumber (preferred) ---
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                for line in text.splitlines():
                    line = _clean(line)
                    if not line:
                        continue
                    # Heuristic: short ALL-CAPS or numbered lines are headings
                    is_heading = (
                        (line.isupper() and len(line) < 80) or
                        re.match(r"^(\d+\.)+\s+\w", line)
                    )
                    blocks.append({
                        "text": line,
                        "section": f"Page {i}",
                        "type": "heading" if is_heading else "paragraph"
                    })
        return blocks
    except ImportError:
        pass

    # --- PyPDF2 fallback ---
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            for line in text.splitlines():
                line = _clean(line)
                if line:
                    blocks.append({"text": line, "section": f"Page {i}", "type": "paragraph"})
        return blocks
    except ImportError:
        log.warning("No PDF library found. Install pdfplumber or PyPDF2.")
        return []


def extract_txt(path: Path) -> list[dict]:
    text = path.read_text(encoding="utf-8", errors="replace")
    blocks = []
    for line in text.splitlines():
        line = _clean(line)
        if not line:
            continue
        blocks.append({"text": line, "section": "Document", "type": "paragraph"})
    return blocks


def extract_md(path: Path) -> list[dict]:
    text = path.read_text(encoding="utf-8", errors="replace")
    blocks = []
    current_section = "Document"
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if heading_match:
            current_section = heading_match.group(2)
            blocks.append({"text": current_section, "section": current_section, "type": "heading"})
        elif stripped.startswith(("- ", "* ", "+ ", "1.")):
            blocks.append({"text": stripped, "section": current_section, "type": "list"})
        else:
            blocks.append({"text": stripped, "section": current_section, "type": "paragraph"})
    return blocks


def extract_csv(path: Path, rows_per_chunk: int = 1) -> list[dict]:
    """
    Each logical row-group becomes a block.
    rows_per_chunk=1  → one chunk per row
    rows_per_chunk>1  → batch rows together (useful for short rows)
    """
    blocks = []
    with path.open(encoding="utf-8-sig", errors="replace") as f:
        reader = csv.DictReader(f)
        headers = reader.fieldnames or []
        buffer: list[str] = []
        for row in reader:
            line = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
            buffer.append(line)
            if len(buffer) >= rows_per_chunk:
                blocks.append({
                    "text": "\n".join(buffer),
                    "section": "Table",
                    "type": "row"
                })
                buffer = []
        if buffer:
            blocks.append({"text": "\n".join(buffer), "section": "Table", "type": "row"})
    return blocks


def extract_json(path: Path, text_field: str | None = None) -> list[dict]:
    """
    - If JSON is a list of objects: each object → one block.
    - If JSON is a dict: recurse into values and emit key-value pairs.
    - text_field: if set, only extract that key from each record.
    """
    raw = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    blocks = []

    def _dict_to_text(d: dict) -> str:
        return "\n".join(f"{k}: {v}" for k, v in d.items() if isinstance(v, (str, int, float, bool)))

    if isinstance(raw, list):
        for i, item in enumerate(raw):
            if isinstance(item, dict):
                text = item.get(text_field, _dict_to_text(item)) if text_field else _dict_to_text(item)
                blocks.append({"text": _clean(str(text)), "section": f"Record {i+1}", "type": "record"})
            else:
                blocks.append({"text": _clean(str(item)), "section": f"Record {i+1}", "type": "record"})
    elif isinstance(raw, dict):
        for key, val in raw.items():
            text = json.dumps(val, ensure_ascii=False) if isinstance(val, (dict, list)) else str(val)
            blocks.append({"text": f"{key}: {_clean(text)}", "section": key, "type": "record"})
    else:
        blocks.append({"text": _clean(str(raw)), "section": "Document", "type": "paragraph"})

    return blocks


def extract_xlsx(path: Path) -> list[dict]:
    try:
        import openpyxl
    except ImportError:
        log.error("openpyxl not installed: pip install openpyxl")
        return []
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    blocks = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_data = []
        headers = None
        for row in ws.iter_rows(values_only=True):
            values = [str(c) if c is not None else "" for c in row]
            if not any(values):
                continue
            if headers is None:
                headers = values
            else:
                line = " | ".join(f"{h}: {v}" for h, v in zip(headers, values) if v)
                if line:
                    rows_data.append(line)
        if rows_data:
            blocks.append({
                "text": "\n".join(rows_data),
                "section": sheet_name,
                "type": "table"
            })
    return blocks


def extract_html(path: Path) -> list[dict]:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        log.warning("beautifulsoup4 not installed; falling back to regex strip.")
        text = re.sub(r"<[^>]+>", " ", path.read_text(encoding="utf-8", errors="replace"))
        return [{"text": _clean(text), "section": "Document", "type": "paragraph"}]

    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="replace"), "html.parser")
    blocks = []
    current_section = "Document"
    for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th"]):
        text = _clean(tag.get_text(separator=" "))
        if not text:
            continue
        if tag.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            current_section = text
            blocks.append({"text": text, "section": current_section, "type": "heading"})
        elif tag.name == "li":
            blocks.append({"text": f"• {text}", "section": current_section, "type": "list"})
        else:
            blocks.append({"text": text, "section": current_section, "type": "paragraph"})
    return blocks


def extract_doc(path: Path) -> list[dict]:
    """
    Extracts text from old .doc files.
    Priority 1: Windows Native (using pywin32 + MS Word)
    Priority 2: Unstructured (requires LibreOffice)
    """
    # --- Priority 1: Windows + MS Word ---
    if os.name == 'nt':
        try:
            import win32com.client
            # Word requires absolute paths
            abs_path = str(path.absolute())
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(abs_path)
                text = doc.Content.Text
                doc.Close()
                
                blocks = []
                for line in text.splitlines():
                    line = _clean(line)
                    if line:
                        blocks.append({"text": line, "section": "Document", "type": "paragraph"})
                return blocks
            finally:
                # word.Quit() # Uncomment if you want to close Word after each file
                pass
        except Exception as e:
            log.debug(f"Win32 Word extraction failed for {path.name}: {e}")

    # --- Priority 2: Unstructured ---
    try:
        from unstructured.partition.doc import partition_doc
        elements = partition_doc(filename=str(path))
        blocks = []
        current_section = "Document"
        for el in elements:
            text = _clean(el.text)
            if not text: continue
            el_type = str(type(el)).lower()
            if "title" in el_type:
                current_section = text
                blocks.append({"text": text, "section": current_section, "type": "heading"})
            else:
                blocks.append({"text": text, "section": current_section, "type": "paragraph"})
        return blocks
    except Exception as e:
        log.error(f"Error extracting .doc file {path.name}: {e}. (Tried Word and Unstructured).")
        return []


# ──────────────────────────────────────────────────────────────────────────────
# DISPATCHER
# ──────────────────────────────────────────────────────────────────────────────

EXTRACTOR_MAP = {
    ".docx": extract_docx,
    ".doc":  extract_doc,
    ".pdf":  extract_pdf,
    ".txt":  extract_txt,
    ".md":   extract_md,
    ".csv":  extract_csv,
    ".json": extract_json,
    ".xlsx": extract_xlsx,
    ".html": extract_html,
    ".htm":  extract_html,
}

AUTO_STRATEGY = {
    ".docx": "SECTION",
    ".doc":  "SECTION",
    ".pdf":  "SECTION",
    ".txt":  "PARAGRAPH",
    ".md":   "SECTION",
    ".csv":  "ROW",
    ".json": "ROW",
    ".xlsx": "ROW",
    ".html": "SECTION",
    ".htm":  "SECTION",
}


# ──────────────────────────────────────────────────────────────────────────────
# CHUNKING STRATEGIES
# ──────────────────────────────────────────────────────────────────────────────

def strategy_section(blocks: list[dict], chunk_size: int, overlap: int) -> list[str]:
    """
    Group consecutive non-heading blocks under the same heading into one chunk.
    If a group exceeds chunk_size, apply fixed windows on top.
    """
    sections: dict[str, list[str]] = {}
    order: list[str] = []
    for b in blocks:
        sec = b["section"]
        if sec not in sections:
            sections[sec] = []
            order.append(sec)
        if b["type"] != "heading":
            sections[sec].append(b["text"])

    chunks = []
    for sec in order:
        body = "\n".join(sections[sec])
        header = f"[{sec}]\n"
        full = header + body
        if len(full) <= chunk_size:
            chunks.append(full)
        else:
            for window in _fixed_windows(body, chunk_size, overlap):
                chunks.append(header + window)
    return chunks


def strategy_paragraph(blocks: list[dict], chunk_size: int, overlap: int) -> list[str]:
    """Each paragraph/list item is a chunk; merge short ones, split long ones."""
    chunks = []
    buffer = ""
    for b in blocks:
        text = b["text"]
        if len(buffer) + len(text) + 1 <= chunk_size:
            buffer = (buffer + "\n" + text).strip()
        else:
            if buffer:
                chunks.append(buffer)
            buffer = text
    if buffer:
        chunks.append(buffer)
    return chunks


def strategy_fixed(blocks: list[dict], chunk_size: int, overlap: int) -> list[str]:
    """Concatenate all text then apply sliding windows."""
    full_text = "\n".join(b["text"] for b in blocks)
    return _fixed_windows(full_text, chunk_size, overlap)


def strategy_row(blocks: list[dict], *_) -> list[str]:
    """One chunk per block (used for CSV/JSON/Excel rows)."""
    return [b["text"] for b in blocks]


STRATEGY_FN = {
    "SECTION":   strategy_section,
    "PARAGRAPH": strategy_paragraph,
    "FIXED":     strategy_fixed,
    "ROW":       strategy_row,
}


# ──────────────────────────────────────────────────────────────────────────────
# CORE CHUNKER
# ──────────────────────────────────────────────────────────────────────────────

class GenericChunker:
    def __init__(
        self,
        strategy: str = "AUTO",
        chunk_size: int = 800,
        overlap: int = 100,
        rows_per_chunk: int = 1,
        json_text_field: str | None = None,
        min_chunk_len: int = 30,
    ):
        self.strategy = strategy.upper()
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.rows_per_chunk = rows_per_chunk
        self.json_text_field = json_text_field
        self.min_chunk_len = min_chunk_len

    def _effective_strategy(self, suffix: str) -> str:
        if self.strategy == "AUTO":
            return AUTO_STRATEGY.get(suffix, "PARAGRAPH")
        return self.strategy

    def _extract(self, path: Path) -> list[dict]:
        suffix = path.suffix.lower()
        extractor = EXTRACTOR_MAP.get(suffix)
        if not extractor:
            log.warning(f"Unsupported file type: {suffix} ({path.name})")
            return []

        # Pass extra kwargs for csv/json extractors
        if suffix == ".csv":
            return extractor(path, rows_per_chunk=self.rows_per_chunk)
        if suffix == ".json":
            return extractor(path, text_field=self.json_text_field)
        return extractor(path)

    def chunk_file(self, path: Path) -> list[Chunk]:
        path = Path(path)
        if not path.exists():
            log.error(f"File not found: {path}")
            return []

        suffix = path.suffix.lower()
        blocks = self._extract(path)
        if not blocks:
            log.warning(f"No content extracted from {path.name}")
            return []

        eff_strategy = self._effective_strategy(suffix)
        strategy_fn = STRATEGY_FN.get(eff_strategy, strategy_paragraph)
        raw_chunks = strategy_fn(blocks, self.chunk_size, self.overlap)

        chunks = []
        for idx, text in enumerate(raw_chunks):
            text = text.strip()
            if len(text) < self.min_chunk_len:
                continue

            # Build metadata
            section = "Document"
            for b in blocks:
                if b["text"] in text:
                    section = b.get("section", "Document")
                    break

            meta = {
                "source":    str(path),
                "file":      path.name,
                "doc_name":  _slug(path),
                "extension": suffix.lstrip("."),
                "strategy":  eff_strategy,
                "section":   section,
                "chunk_id":  idx,
            }
            chunks.append(Chunk(text=text, metadata=meta))

        log.info(f"  {path.name} → {len(chunks)} chunks (strategy={eff_strategy})")
        return chunks

    def chunk_directory(
        self,
        directory: Path,
        recursive: bool = False,
        include_exts: list[str] | None = None,
        exclude_exts: list[str] | None = None,
    ) -> list[Chunk]:
        directory = Path(directory)
        pattern = "**/*" if recursive else "*"
        all_chunks: list[Chunk] = []

        supported = set(EXTRACTOR_MAP.keys())
        include = {f".{e.lstrip('.')}" for e in include_exts} if include_exts else supported
        exclude = {f".{e.lstrip('.')}" for e in exclude_exts} if exclude_exts else set()
        target_exts = (include & supported) - exclude

        files = [p for p in directory.glob(pattern) if p.is_file() and p.suffix.lower() in target_exts]
        log.info(f"Found {len(files)} file(s) in {directory}")

        for f in sorted(files):
            all_chunks.extend(self.chunk_file(f))

        return all_chunks


# ──────────────────────────────────────────────────────────────────────────────
# OUTPUT
# ──────────────────────────────────────────────────────────────────────────────

def save_jsonl(chunks: list[Chunk], output_path: str | Path) -> None:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8") as f:
        for chunk in chunks:
            f.write(chunk.to_jsonl() + "\n")
    log.info(f"Saved {len(chunks)} chunks → {output_path}")


def build_faiss_index(
    chunks: list[Chunk],
    vectorstore_dir: str | Path,
    index_name: str = "index",
    model: str = "BAAI/bge-m3",
    device: str = "cpu",
) -> None:
    try:
        from langchain_community.vectorstores import FAISS
        from langchain_huggingface import HuggingFaceEmbeddings
        from langchain_core.documents import Document as LCDoc
    except ImportError:
        log.error("LangChain not installed: pip install langchain langchain-community langchain-huggingface sentence-transformers faiss-cpu")
        return

    log.info(f"Loading embedding model: {model}")
    embeddings = HuggingFaceEmbeddings(
        model_name=model,
        model_kwargs={"device": device},
        encode_kwargs={"normalize_embeddings": True},
    )

    lc_docs = [LCDoc(page_content=c.text, metadata=c.metadata) for c in chunks]

    log.info(f"Building FAISS index over {len(lc_docs)} chunks...")
    vectorstore = FAISS.from_documents(lc_docs, embeddings)

    out = Path(vectorstore_dir)
    out.mkdir(parents=True, exist_ok=True)
    vectorstore.save_local(str(out), index_name=index_name)
    log.info(f"FAISS index saved → {out}/{index_name}.faiss")


# ──────────────────────────────────────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────────────────────────────────────

def build_cli() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="generic_chunker",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        description=textwrap.dedent("""\
            Generic document chunker for embedding pipelines.

            Examples:
              # Chunk all files in a folder, save JSONL
              python generic_chunker.py --input ./docs --output chunks.jsonl

              # Also build a FAISS vector store
              python generic_chunker.py --input ./docs --output chunks.jsonl --embed

              # Chunk a single PDF with fixed-size windows
              python generic_chunker.py --input report.pdf --strategy FIXED --chunk-size 512 --overlap 64

              # Chunk CSV with 3 rows per chunk
              python generic_chunker.py --input data.csv --strategy ROW --rows-per-chunk 3

              # Only include .docx and .pdf files recursively
              python generic_chunker.py --input ./docs --recursive --include docx pdf
        """),
    )
    p.add_argument("--input",          required=True,          help="File or directory to chunk")
    p.add_argument("--output",         default="chunks.jsonl", help="Output JSONL path")
    p.add_argument("--strategy",       default="AUTO",
                   choices=["AUTO", "SECTION", "PARAGRAPH", "FIXED", "ROW"],
                   help="Chunking strategy (default: AUTO)")
    p.add_argument("--chunk-size",     type=int, default=800,  help="Max chars per chunk (default: 800)")
    p.add_argument("--overlap",        type=int, default=100,  help="Overlap chars for FIXED strategy (default: 100)")
    p.add_argument("--rows-per-chunk", type=int, default=1,    help="Rows per chunk for CSV/Excel (default: 1)")
    p.add_argument("--json-field",     default=None,           help="JSON key to use as text field")
    p.add_argument("--min-len",        type=int, default=30,   help="Minimum chunk length to keep (default: 30)")
    p.add_argument("--recursive",      action="store_true",    help="Recurse into subdirectories")
    p.add_argument("--include",        nargs="+",              help="File extensions to include (e.g. docx pdf)")
    p.add_argument("--exclude",        nargs="+",              help="File extensions to exclude")
    p.add_argument("--embed",          action="store_true",    help="Build FAISS vector store after chunking")
    p.add_argument("--model",          default="BAAI/bge-m3",  help="Embedding model (default: BAAI/bge-m3)")
    p.add_argument("--device",         default="cpu",          help="Device for embeddings: cpu or cuda")
    p.add_argument("--vectorstore",    default="vectorstore",  help="FAISS output directory")
    p.add_argument("--index-name",     default="index",        help="FAISS index name")
    return p


def main(argv: list[str] | None = None) -> None:
    parser = build_cli()
    args = parser.parse_args(argv)

    chunker = GenericChunker(
        strategy=args.strategy,
        chunk_size=args.chunk_size,
        overlap=args.overlap,
        rows_per_chunk=args.rows_per_chunk,
        json_text_field=args.json_field,
        min_chunk_len=args.min_len,
    )

    input_path = Path(args.input)
    if input_path.is_dir():
        chunks = chunker.chunk_directory(
            input_path,
            recursive=args.recursive,
            include_exts=args.include,
            exclude_exts=args.exclude,
        )
    else:
        chunks = chunker.chunk_file(input_path)

    if not chunks:
        log.warning("No chunks produced. Check input files and options.")
        sys.exit(1)

    save_jsonl(chunks, args.output)

    if args.embed:
        build_faiss_index(
            chunks,
            vectorstore_dir=args.vectorstore,
            index_name=args.index_name,
            model=args.model,
            device=args.device,
        )

    log.info(f"Done. Total chunks: {len(chunks)}")


# ──────────────────────────────────────────────────────────────────────────────
# PYTHON API USAGE EXAMPLE
# ──────────────────────────────────────────────────────────────────────────────
#
# from generic_chunker import GenericChunker, save_jsonl, build_faiss_index
#
# chunker = GenericChunker(strategy="AUTO", chunk_size=800, overlap=100)
#
# # Single file
# chunks = chunker.chunk_file("report.pdf")
#
# # Entire folder (only docx + pdf, recursive)
# chunks = chunker.chunk_directory("./docs", recursive=True, include_exts=["docx", "pdf"])
#
# save_jsonl(chunks, "output.jsonl")
# build_faiss_index(chunks, "vectorstore", model="BAAI/bge-m3")
#
# ──────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    main()

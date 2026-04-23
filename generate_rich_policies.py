import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ==============================
# DOC WRITER (ENTERPRISE STYLE)
# ==============================
def write_enterprise_doc(filepath, title, sections):
    doc = Document()

    # HEADER
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.text = "CHILLI CORP | INTERNAL POLICY HANDBOOK | CONFIDENTIAL"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # TITLE PAGE
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Document Type: Corporate Policy Handbook\n"
        "Classification: INTERNAL USE ONLY\n"
        "Version: 2.0\n"
        "Approval Authority: Executive Board\n"
        "Legal Framework: Vietnam Labor Law 2019 + Internal Governance Rules\n"
        "Document Status: ACTIVE"
    )

    doc.add_page_break()

    # CONTENT
    for i, sec in enumerate(sections, 1):
        doc.add_heading(f"{i}. {sec['title']}", level=1)

        for block in sec["content"]:
            if isinstance(block, str):
                doc.add_paragraph(block)

            elif isinstance(block, dict) and block["type"] == "table":
                table = doc.add_table(rows=1, cols=len(block["headers"]))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                hdr = table.rows[0].cells
                for j, h in enumerate(block["headers"]):
                    hdr[j].text = h

                for row in block["rows"]:
                    r = table.add_row().cells
                    for j, v in enumerate(row):
                        r[j].text = v

    doc.save(filepath)
    print(f"✔ Generated: {filepath}")


# ==============================
# ENTERPRISE CONTENT ENGINE
# ==============================
def policy_block(policy):
    return [
        {
            "title": "Legal Basis & Governance Framework",
            "content": [
                f"This {policy} policy is established in accordance with Vietnamese Labor Law and internal corporate governance standards.",
                "It defines mandatory rules applicable to all employees regardless of position or department.",
                "Any deviation from this policy must be formally approved by authorized management.",
                "The policy ensures legal compliance, operational stability, and risk mitigation."
            ]
        },
        {
            "title": "Scope of Application",
            "content": [
                "This policy applies to all full-time employees, contractors, interns, and temporary staff.",
                "It is mandatory across all departments including Engineering, HR, Finance, and Operations.",
                "Third-party vendors must comply when operating within company systems."
            ]
        },
        {
            "title": "Definitions & Terminology",
            "content": [
                "- Employee: Any individual under contract with the company.",
                "- Manager: Person responsible for team supervision and approval.",
                "- HR Department: Authority for policy enforcement and compliance.",
                "- Violation: Any action inconsistent with policy rules."
            ]
        },
        {
            "title": "Operational Procedures",
            "content": [
                "Step 1: Request submission via internal system.",
                "Step 2: Manager review and validation.",
                "Step 3: Department-level approval.",
                "Step 4: HR compliance verification.",
                "Step 5: System logging and archival.",
                "All procedures must be executed in sequence without bypass."
            ]
        },
        {
            "title": "Roles & Responsibilities",
            "content": [
                "- Employees: Must comply with all policy requirements.",
                "- Managers: Responsible for approval and enforcement.",
                "- HR: Ensures policy execution and monitoring.",
                "- IT: Maintains system integrity and access control."
            ]
        },
        {
            "title": "Compliance, Risk & Audit Control",
            "content": [
                "Non-compliance may result in disciplinary action including termination.",
                "The company conducts quarterly internal audits and annual external audits.",
                "All actions are logged for traceability and accountability.",
                "Risk categories include operational risk, legal risk, and data security risk.",
                "Mitigation strategies are enforced through automated monitoring systems."
            ]
        },
        {
            "title": "Exception Handling Mechanism",
            "content": [
                "Exceptions are only granted under critical or emergency conditions.",
                "Approval must come from at least 2 management levels.",
                "All exceptions must be documented and stored in audit logs.",
                "Repeated exception requests will trigger compliance review."
            ]
        },
        {
            "title": "Risk Classification Table",
            "content": [
                {
                    "type": "table",
                    "headers": ["Risk Type", "Description", "Severity", "Mitigation"],
                    "rows": [
                        ["Operational", "Workflow disruption", "Medium", "Process automation"],
                        ["Legal", "Regulatory violation", "High", "Legal review"],
                        ["Security", "Data breach", "Critical", "Access control + encryption"]
                    ]
                }
            ]
        },
        {
            "title": "Real-world Application Scenarios",
            "content": [
                "Scenario 1: Employee follows full process → Approved and logged.",
                "Scenario 2: Employee skips approval → System auto-rejects.",
                "Scenario 3: Emergency request → Fast-track approval applied.",
                "Scenario 4: Repeated violation → HR escalation triggered."
            ]
        }
    ]


# ==============================
# FULL 20 FILES (FIXED STRUCTURE)
# ==============================
CONTENT_LIBRARY = {
    "Leave": [{"title": "Leave Management System", "content": ["Defines employee leave entitlement and approval system."]}],
    "Working_Hours": [{"title": "Work Time Policy", "content": ["Defines standard working hours and overtime rules."]}],
    "Remote_Work": [{"title": "Remote Work Framework", "content": ["Defines WFH eligibility and security rules."]}],
    "Code_of_Conduct": [{"title": "Ethical Conduct Policy", "content": ["Defines workplace behavior and ethics standards."]}],
    "IT_Usage": [{"title": "IT Resource Policy", "content": ["Defines proper usage of company IT systems."]}],
    "Confidentiality": [{"title": "Data Confidentiality Policy", "content": ["Protects internal and client information."]}],
    "Onboarding": [{"title": "Employee Onboarding System", "content": ["Defines onboarding workflow and training."]}],
    "Performance": [{"title": "Performance Management", "content": ["Defines KPI-based evaluation system."]}],
    "Expense": [{"title": "Expense Control Policy", "content": ["Defines reimbursement and financial control rules."]}],
    "Travel": [{"title": "Business Travel Policy", "content": ["Defines travel approval and allowance rules."]}],
    "Dress_Code": [{"title": "Professional Dress Code", "content": ["Defines corporate appearance standards."]}],
    "Attendance": [{"title": "Attendance Control System", "content": ["Defines check-in/out and punctuality enforcement."]}],
    "Data_Security": [{"title": "Information Security Policy", "content": ["Defines data protection and access control rules."]}],
    "Training": [{"title": "Training & Development", "content": ["Defines employee learning framework."]}],
    "Promotion": [{"title": "Career Advancement Policy", "content": ["Defines promotion criteria and process."]}],
    "Termination": [{"title": "Employment Termination Policy", "content": ["Defines exit procedures and compliance rules."]}],
    "Health": [{"title": "Health & Safety Policy", "content": ["Ensures workplace safety compliance."]}],
    "Work_From_Home": [{"title": "WFH Policy", "content": ["Defines remote working conditions."]}],
    "Email": [{"title": "Email Usage Policy", "content": ["Defines corporate communication rules."]}],
    "Meeting": [{"title": "Meeting Governance Policy", "content": ["Defines meeting discipline and documentation rules."]}]
}


# ==============================
# GENERATOR ENGINE
# ==============================
def generate_all():
    base_dir = r"d:\All Project\Chilli_Assignment\docs\internal_policies"
    os.makedirs(base_dir, exist_ok=True)

    keys = list(CONTENT_LIBRARY.keys())

    for i in range(1, 21):
        key = keys[i - 1]

        filename = f"{i:02d}_{key}_Policy.docx"
        filepath = os.path.join(base_dir, filename)

        base_sections = [dict(s) for s in CONTENT_LIBRARY[key]]
        full_sections = base_sections + policy_block(key)

        write_enterprise_doc(
            filepath,
            f"{key.upper()} POLICY HANDBOOK",
            full_sections
        )


if __name__ == "__main__":
    generate_all()